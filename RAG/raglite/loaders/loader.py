"""
Document loader module for RAGLite.

Supports loading and chunking of PDF, DOCX, and TXT files.
"""

import os
import logging
import time
from datetime import datetime
from typing import List, Dict, Any, Optional
from pathlib import Path
import tempfile

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

from langchain.text_splitter import RecursiveCharacterTextSplitter

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentLoader:
    """
    A comprehensive document loader that supports PDF, DOCX, and TXT files.
    Includes intelligent text chunking for optimal RAG performance.
    """
    
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        """
        Initialize the document loader.
        
        Args:
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Initialize text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # Check for required dependencies
        self._check_dependencies()
    
    def _check_dependencies(self) -> None:
        """Check if required dependencies are available."""
        if PyPDF2 is None:
            logger.warning("PyPDF2 not installed. PDF loading will not be available.")
        if DocxDocument is None:
            logger.warning("python-docx not installed. DOCX loading will not be available.")
    
    def load_document(self, file_path: str) -> List[Dict[str, Any]]:
        """
        Load a document and return chunked text with metadata.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            List of dictionaries containing chunked text and metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")
        
        file_extension = Path(file_path).suffix.lower()
        
        # Extract text based on file type
        if file_extension == '.pdf':
            text = self._load_pdf(file_path)
        elif file_extension == '.docx':
            text = self._load_docx(file_path)
        elif file_extension == '.txt':
            text = self._load_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create documents with metadata
        documents = []
        upload_time = datetime.now().isoformat()
        
        for i, chunk in enumerate(chunks):
            # Determine page numbers for this chunk
            page_numbers = self._get_page_numbers_for_chunk(chunk, file_extension)
            
            doc = {
                'text': chunk,
                'metadata': {
                    'source': file_path,
                    'filename': Path(file_path).name,
                    'chunk_id': i,
                    'total_chunks': len(chunks),
                    'file_type': file_extension,
                    'chunk_size': len(chunk),
                    'upload_time': upload_time,
                    'page_numbers': page_numbers
                }
            }
            documents.append(doc)
        
        logger.info(f"Loaded {len(documents)} chunks from {Path(file_path).name}")
        return documents
    
    def _get_page_numbers_for_chunk(self, chunk_text: str, file_type: str) -> List[int]:
        """
        Extract page numbers from a chunk of text.
        
        Args:
            chunk_text: The text chunk
            file_type: Type of file (for PDF files)
            
        Returns:
            List of page numbers found in the chunk
        """
        import re
        
        page_numbers = []
        
        if file_type == '.pdf' and hasattr(self, '_page_info'):
            # For PDF files, check which pages are included in this chunk
            for page_info in self._page_info:
                if page_info['text'] in chunk_text:
                    page_numbers.append(page_info['page_number'])
        else:
            # For other file types, try to extract page numbers from text
            page_patterns = [
                r'--- Page (\d+) ---',
                r'Page (\d+)',
                r'page (\d+)',
                r'PAGE (\d+)'
            ]
            
            for pattern in page_patterns:
                matches = re.findall(pattern, chunk_text)
                for match in matches:
                    try:
                        page_num = int(match)
                        if page_num not in page_numbers:
                            page_numbers.append(page_num)
                    except (ValueError, IndexError):
                        continue
        
        return sorted(page_numbers)
    
    def _load_pdf(self, file_path: str) -> str:
        """Load text from a PDF file."""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 is required for PDF loading. Install with: pip install PyPDF2")
        
        text = ""
        page_info = []  # Store page information for metadata
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():  # Only add non-empty pages
                            page_number = page_num + 1
                            text += f"\n--- Page {page_number} ---\n{page_text}\n"
                            page_info.append({
                                'page_number': page_number,
                                'text': page_text,
                                'start_pos': len(text) - len(page_text) - len(f"\n--- Page {page_number} ---\n")
                            })
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
        
        if not text.strip():
            raise ValueError(f"No text could be extracted from PDF: {file_path}")
        
        # Store page info for later use in chunking
        self._page_info = page_info
        return text
    
    def _load_docx(self, file_path: str) -> str:
        """Load text from a DOCX file."""
        if DocxDocument is None:
            raise ImportError("python-docx is required for DOCX loading. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
        
        if not text.strip():
            raise ValueError(f"No text could be extracted from DOCX: {file_path}")
        
        return text
    
    def _load_txt(self, file_path: str) -> str:
        """Load text from a TXT file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        except UnicodeDecodeError:
            # Try with different encoding if UTF-8 fails
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
            except Exception as e:
                logger.error(f"Error loading TXT {file_path}: {e}")
                raise
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise
        
        if not text.strip():
            raise ValueError(f"Empty text file: {file_path}")
        
        return text
    
    def load_from_bytes(self, file_bytes: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Load a document from bytes (useful for uploaded files).
        
        Args:
            file_bytes: Raw bytes of the file
            filename: Original filename with extension
            
        Returns:
            List of dictionaries containing chunked text and metadata
        """
        file_extension = Path(filename).suffix.lower()
        
        # Create temporary file
        with tempfile.NamedTemporaryFile(suffix=file_extension, delete=False) as temp_file:
            temp_file.write(file_bytes)
            temp_path = temp_file.name
        
        try:
            # Load document from temporary file
            documents = self.load_document(temp_path)
            
            # Update metadata to reflect original filename
            for doc in documents:
                doc['metadata']['filename'] = filename
                doc['metadata']['source'] = filename
                # Keep page_numbers if they were extracted
            
            return documents
        finally:
            # Clean up temporary file
            try:
                os.unlink(temp_path)
            except Exception as e:
                logger.warning(f"Could not delete temporary file {temp_path}: {e}")
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported file formats."""
        formats = ['.txt']
        if PyPDF2 is not None:
            formats.append('.pdf')
        if DocxDocument is not None:
            formats.append('.docx')
        return formats 